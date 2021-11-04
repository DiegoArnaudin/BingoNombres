import random
import docx

nombres = []

print('Ingrese Nombres (Vacio para terminar)')
nombre = input()
while nombre != '':
    nombres.append(nombre)    
    nombre = input()

print('Ingrese tamaño de carton:')
filas = int( input('Filas: ') )
columnas = int( input('Columnas: ') )

cartones = []
for j in range(len(nombres)):
    while True:
        carton = []
        for i in range(filas*columnas-1):
            while True:
                nombre = nombres[random.randint(0,len(nombres)-1)]
                if nombre not in carton:
                    carton.append(nombre)
                    break;
        if carton not in cartones:
            cartones.append(carton)
            break;

doc = docx.Document()

for carton in cartones:
    table = doc.add_table(filas,columnas)

    r = 0; c = 0;
    for nombre in carton:
        cell = table.cell(r,c)
        c=c+1
        if c>=columnas:
            c=0;
            r=r+1;
            
        cell.text = nombre

    doc.add_paragraph()

doc.save(r"BINGO.docx")

